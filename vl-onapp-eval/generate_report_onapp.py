"""
generate_report_onapp.py — OnApp VL 评测报告生成器

输出:
  charts/bar_overall.png       — 5模型总分排名（条件A基线）
  charts/heatmap_dims.png      — 12维度×5模型热图（条件A）
  charts/radar_dims.png        — 5模型雷达图（条件A）
  charts/ablation_delta.png    — 消融实验：A→B/C/D的分数变化
  summary_onapp.docx           — 总结报告

用法:
  python generate_report_onapp.py
  python generate_report_onapp.py --run run_20260320_1000
"""

import os, sys, json, argparse
from pathlib import Path

import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm

from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

if sys.stdout.encoding != "utf-8":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

EVAL_DIR     = Path(__file__).parent
RESULTS_BASE = EVAL_DIR / "results"
CHARTS_DIR   = EVAL_DIR / "charts"

DIMS = [
    "视觉识别准确率", "幻觉控制率", "数值读取精度", "输出时序合规",
    "安全声明合规",   "边界克制合规", "数据引用质量", "端侧数据优先性",
    "工具调用时机准确性", "工具调用结果整合度", "任务完成度", "图像与上下文一致性",
]

DIMS_SHORT = [
    "视觉识别", "幻觉控制", "数值精度", "时序合规",
    "安全声明", "边界克制", "引用质量", "端侧优先",
    "工具时机", "工具整合", "任务完成", "图文一致",
]

ABLATION_DIMS = {
    "B": ["视觉识别准确率", "数据引用质量", "图像与上下文一致性"],
    "C": ["数值读取精度", "数据引用质量", "任务完成度"],
    "D": ["数据引用质量", "任务完成度"],
}

MODEL_COLORS = ["#4f86c6", "#e05555", "#55a855", "#e0a500", "#8855cc"]

# ── Font ──────────────────────────────────────────────────────────────────────
_CN: dict = {}
def cn(size=10):
    if size in _CN:
        return _CN[size]
    for name in ["Microsoft YaHei", "SimHei", "SimSun", "Noto Sans CJK SC"]:
        try:
            fp = fm.findfont(fm.FontProperties(family=name), fallback_to_default=False)
            if fp and "DejaVu" not in fp:
                p = fm.FontProperties(family=name, size=size)
                _CN[size] = p
                return p
        except Exception:
            pass
    p = fm.FontProperties(size=size)
    _CN[size] = p
    return p


# ── Data loading ──────────────────────────────────────────────────────────────
def load_results(results_dir: Path) -> dict:
    """Returns {model_name: {condition: [results]}}"""
    out = {}
    for json_path in sorted(results_dir.glob("*_results.json")):
        data = json.loads(json_path.read_text("utf-8"))
        name = data["model"]["name"]
        out[name] = {}
        for cond, results in data.get("conditions", {}).items():
            out[name][cond] = results
    return out


def compute_dim_avgs(results: list, dims: list) -> dict[str, float]:
    """Compute average score per dimension from judged results."""
    acc: dict[str, list] = {d: [] for d in dims}
    for r in results:
        j = r.get("judgment")
        if not j:
            continue
        for d in dims:
            score = j.get("scores", {}).get(d, 0)
            if score > 0:
                acc[d].append(score)
    return {d: (sum(v) / len(v) if v else 0.0) for d, v in acc.items()}


def compute_overall(dim_avgs: dict) -> float:
    vals = [v for v in dim_avgs.values() if v > 0]
    return sum(vals) / len(vals) if vals else 0.0


# ── Chart 1: Bar overall scores (condition A) ─────────────────────────────────
def chart_bar_overall(all_results: dict, output: Path):
    model_names = list(all_results.keys())
    scores = []
    for name in model_names:
        results_A = all_results[name].get("A", [])
        avgs = compute_dim_avgs(results_A, DIMS)
        scores.append(compute_overall(avgs))

    # Sort descending
    pairs = sorted(zip(model_names, scores), key=lambda x: -x[1])
    names_sorted = [p[0] for p in pairs]
    scores_sorted = [p[1] for p in pairs]

    fig, ax = plt.subplots(figsize=(10, 5))
    colors = [MODEL_COLORS[model_names.index(n) % len(MODEL_COLORS)] for n in names_sorted]
    bars = ax.barh(names_sorted[::-1], scores_sorted[::-1], color=colors[::-1], height=0.5)

    for bar, score in zip(bars, scores_sorted[::-1]):
        ax.text(bar.get_width() + 0.05, bar.get_y() + bar.get_height() / 2,
                f"{score:.2f}", va="center", fontproperties=cn(10))

    ax.set_xlim(0, 10.5)
    ax.set_xlabel("综合均分（1-10）", fontproperties=cn(11))
    ax.set_title("OnApp VL 评测 — 总分排名（条件A基线）", fontproperties=cn(13))
    for label in ax.get_yticklabels():
        label.set_fontproperties(cn(10))
    ax.grid(axis="x", alpha=0.3)
    plt.tight_layout()
    plt.savefig(str(output), dpi=150, bbox_inches="tight")
    plt.close()
    print(f"[chart] bar_overall → {output}")


# ── Chart 2: Heatmap (dims × models, condition A) ────────────────────────────
def chart_heatmap(all_results: dict, output: Path):
    model_names = list(all_results.keys())
    matrix = []
    for name in model_names:
        results_A = all_results[name].get("A", [])
        avgs = compute_dim_avgs(results_A, DIMS)
        matrix.append([avgs.get(d, 0) for d in DIMS])

    data = np.array(matrix)  # shape: (models, dims)

    fig, ax = plt.subplots(figsize=(14, 5))
    im = ax.imshow(data, cmap="RdYlGn", aspect="auto", vmin=1, vmax=10)

    ax.set_xticks(range(len(DIMS_SHORT)))
    ax.set_xticklabels(DIMS_SHORT, fontproperties=cn(9))
    ax.set_yticks(range(len(model_names)))
    ax.set_yticklabels(model_names, fontproperties=cn(10))

    for i in range(len(model_names)):
        for j in range(len(DIMS)):
            val = data[i, j]
            color = "white" if val < 4 or val > 8 else "black"
            ax.text(j, i, f"{val:.1f}", ha="center", va="center",
                    color=color, fontproperties=cn(8))

    plt.colorbar(im, ax=ax, label="分数")
    ax.set_title("各维度得分热图（条件A基线）", fontproperties=cn(13))
    plt.setp(ax.get_xticklabels(), rotation=30, ha="right")
    plt.tight_layout()
    plt.savefig(str(output), dpi=150, bbox_inches="tight")
    plt.close()
    print(f"[chart] heatmap → {output}")


# ── Chart 3: Radar (condition A) ─────────────────────────────────────────────
def chart_radar(all_results: dict, output: Path):
    model_names = list(all_results.keys())
    N = len(DIMS_SHORT)
    angles = [n / float(N) * 2 * np.pi for n in range(N)]
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(polar=True))

    for idx, name in enumerate(model_names):
        results_A = all_results[name].get("A", [])
        avgs = compute_dim_avgs(results_A, DIMS)
        values = [avgs.get(d, 0) for d in DIMS]
        values += values[:1]
        color = MODEL_COLORS[idx % len(MODEL_COLORS)]
        ax.plot(angles, values, "o-", linewidth=1.5, color=color, label=name)
        ax.fill(angles, values, alpha=0.1, color=color)

    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(DIMS_SHORT, fontproperties=cn(9))
    ax.set_ylim(0, 10)
    ax.set_yticks([2, 4, 6, 8, 10])
    ax.set_yticklabels(["2", "4", "6", "8", "10"], fontproperties=cn(8))
    ax.set_title("多维能力雷达图（条件A基线）", fontproperties=cn(13), pad=20)
    ax.legend(loc="upper right", bbox_to_anchor=(1.3, 1.1),
              prop=cn(9))
    plt.tight_layout()
    plt.savefig(str(output), dpi=150, bbox_inches="tight")
    plt.close()
    print(f"[chart] radar → {output}")


# ── Chart 4: Ablation delta (B/C/D vs A) ─────────────────────────────────────
def chart_ablation(all_results: dict, output: Path):
    model_names = list(all_results.keys())
    ablation_conditions = ["B", "C", "D"]
    cond_labels = {
        "B": "+description",
        "C": "+knowledge",
        "D": "+both",
    }

    # For each condition, compute avg delta on its key dims vs condition A
    model_deltas: dict[str, dict[str, float]] = {}  # model → cond → delta

    for name in model_names:
        model_deltas[name] = {}
        for cond in ablation_conditions:
            dims_to_compare = ABLATION_DIMS.get(cond, [])
            results_A    = all_results[name].get("A", [])
            results_cond = all_results[name].get(cond, [])

            avgs_A    = compute_dim_avgs(results_A, dims_to_compare)
            avgs_cond = compute_dim_avgs(results_cond, dims_to_compare)

            # Average delta across tracked dims
            deltas = []
            for d in dims_to_compare:
                a_val = avgs_A.get(d, 0)
                c_val = avgs_cond.get(d, 0)
                if a_val > 0 and c_val > 0:
                    deltas.append(c_val - a_val)

            model_deltas[name][cond] = sum(deltas) / len(deltas) if deltas else 0.0

    # Plot grouped bar chart
    n_models = len(model_names)
    n_conds  = len(ablation_conditions)
    x = np.arange(n_conds)
    width = 0.8 / n_models

    fig, ax = plt.subplots(figsize=(10, 6))

    for idx, name in enumerate(model_names):
        deltas = [model_deltas[name].get(c, 0) for c in ablation_conditions]
        offset = (idx - n_models / 2 + 0.5) * width
        bars = ax.bar(x + offset, deltas, width, label=name,
                      color=MODEL_COLORS[idx % len(MODEL_COLORS)], alpha=0.85)
        for bar, val in zip(bars, deltas):
            if abs(val) > 0.05:
                ax.text(bar.get_x() + bar.get_width() / 2,
                        bar.get_height() + (0.02 if val >= 0 else -0.08),
                        f"{val:+.2f}", ha="center", va="bottom",
                        fontproperties=cn(8))

    ax.axhline(0, color="black", linewidth=0.8)
    ax.set_xticks(x)
    ax.set_xticklabels([cond_labels[c] for c in ablation_conditions],
                       fontproperties=cn(11))
    ax.set_ylabel("得分变化（vs 条件A基线）", fontproperties=cn(11))
    ax.set_title("消融实验：description/knowledge注入效果", fontproperties=cn(13))
    ax.legend(prop=cn(9))
    ax.grid(axis="y", alpha=0.3)
    plt.tight_layout()
    plt.savefig(str(output), dpi=150, bbox_inches="tight")
    plt.close()
    print(f"[chart] ablation_delta → {output}")


# ── DOCX report ───────────────────────────────────────────────────────────────
def generate_docx(all_results: dict, charts_dir: Path, output: Path):
    doc = Document()
    doc.add_heading("PHA OnApp VL 评测报告", 0)
    doc.add_paragraph("评测条件: A(基线) | B(+description) | C(+knowledge) | D(+both)")
    doc.add_paragraph("评测维度: 12 | 裁判: Claude Code 人工评分 | 测试规模: 5模型 × 4条件 × 48 case")
    doc.add_paragraph("")

    # ── Executive summary (hand-written) ──────────────────────────────────────
    doc.add_heading("总结", 1)

    # Compute key stats for narrative
    model_names = list(all_results.keys())
    rank_rows = []
    for name in model_names:
        results_A = all_results[name].get("A", [])
        avgs_A = compute_dim_avgs(results_A, DIMS)
        overall = compute_overall(avgs_A)
        rank_rows.append((name, overall, avgs_A))
    rank_rows.sort(key=lambda x: -x[1])

    # Compute ablation deltas (dim-specific comparison, same as chart)
    abl_deltas: dict[str, dict[str, float]] = {}
    for name, _, avgs_A in rank_rows:
        abl_deltas[name] = {}
        for cond in ["B", "C", "D"]:
            dims_cmp = ABLATION_DIMS[cond]
            results_cond = all_results[name].get(cond, [])
            avgs_cond = compute_dim_avgs(results_cond, dims_cmp)
            avgs_A_cmp = {d: avgs_A.get(d, 0) for d in dims_cmp}
            deltas = [avgs_cond[d] - avgs_A_cmp[d]
                      for d in dims_cmp if avgs_A_cmp.get(d, 0) > 0 and avgs_cond.get(d, 0) > 0]
            abl_deltas[name][cond] = sum(deltas) / len(deltas) if deltas else 0.0

    # Find best/worst, most/least injection-sensitive
    best_name, best_score, best_avgs = rank_rows[0]
    worst_name, worst_score, worst_avgs = rank_rows[-1]
    most_gain_name = max(model_names, key=lambda n: abl_deltas[n].get("D", 0))
    most_gain_val = abl_deltas[most_gain_name]["D"]
    glm_d_delta = abl_deltas.get("GLM-4.6V", {}).get("D", 0)

    # Find desc-hurts models (B delta < 0)
    desc_hurt = [(n, abl_deltas[n]["B"]) for n in model_names if abl_deltas[n].get("B", 0) < 0]
    desc_hurt.sort(key=lambda x: x[1])

    # 235B hallucination vs vision
    name_235 = next((n for n in model_names if "235" in n), None)
    score_235 = next((s for n, s, _ in rank_rows if n == name_235), 0) if name_235 else 0
    avgs_235 = next((a for n, _, a in rank_rows if n == name_235), {}) if name_235 else {}

    p_exec = doc.add_paragraph()
    p_exec.add_run(
        f"本次评测共5个视觉语言模型，基于30张健康App真实截图和48个查询任务，"
        f"在4种上下文注入条件下系统对比了模型的多维健康助理能力（含视觉识别、幻觉控制、"
        f"工具调用时机、任务完成等12个维度）。\n\n"
    )

    p_exec.add_run("整体格局：").bold = True
    p_exec.add_run(
        f"{best_name}以{best_score:.2f}分位居第一，视觉识别和数值精度尤为突出；"
        f"{worst_name}以{worst_score:.2f}分垫底，工具调用整合（"
        f"{worst_avgs.get('工具调用结果整合度', 0):.1f}）和安全声明合规（"
        f"{worst_avgs.get('安全声明合规', 0):.1f}）是主要拖分项。"
        f"中间三档差距约0.5分，整体呈现kimi→397B→235B→122B的梯队格局。\n\n"
    )

    p_exec.add_run("注入效果分化：").bold = True
    hurt_str = "、".join(f"{n}（{v:+.2f}）" for n, v in desc_hurt)
    p_exec.add_run(
        f"description注入在{hurt_str}等模型上导致视觉相关维度微降，"
        f"说明对视觉能力已强的模型，额外的场景文字描述带来的干扰大于辅助。"
        f"knowledge注入对{most_gain_name}效果最显著"
        f"（整体D条件较A提升{most_gain_val:+.2f}），"
        f"知识注入越多、任务完成度提升越明显，体现知识空白填补效应。"
        f'GLM-4.6V是唯一"注入越多得分越低"的模型（D较A下降{abs(glm_d_delta):.2f}），'
        f"提示其上下文整合能力存在根本缺陷。\n\n"
    )

    if name_235:
        p_exec.add_run("值得关注：").bold = True
        p_exec.add_run(
            f"{name_235}的profile最为特殊——视觉识别全场第二（"
            f"{avgs_235.get('视觉识别准确率', 0):.2f}），"
            f"幻觉控制却全场最低（{avgs_235.get('幻觉控制率', 0):.2f}），"
            f"且存在严重的工具过度调用（非异常case仍频繁触发云API）。"
            f'这体现了大参数模型在健康场景下"会看、但不会守边界"的典型问题。\n\n'
        )

    p_exec.add_run("已知局限：").bold = True
    p_exec.add_run(
        "本次评分由Claude Code（claude-sonnet-4-6）担任唯一裁判，与被评测模型无重叠，"
        "但单一裁判可能对不同厂商的表达风格存在系统性倾向，建议后续引入人工抽检交叉验证；"
        "30张图被多case复用，若裁判对特定图片产生固定误判，误差会跨case放大；"
        "GLM条件C因模型崩溃只有43/48有效sample，结论置信度略低于其他条件。"
    )
    doc.add_paragraph("")

    # Ablation conditions description
    doc.add_heading("评测条件说明（消融设计）", 1)
    p_intro = doc.add_paragraph()
    p_intro.add_run(
        "本评测采用4条件消融设计，系统性验证 description.json 和 knowledge.json 注入对模型表现的影响。"
        "每个模型在全部4个条件下各运行48个case（共192次调用）。\n\n"
    )
    cond_table = doc.add_table(rows=1, cols=4)
    cond_table.style = "Table Grid"
    ch = cond_table.rows[0].cells
    ch[0].text = "条件"
    ch[1].text = "输入内容"
    ch[2].text = "核心评估维度"
    ch[3].text = "设计目的"
    cond_data = [
        ("A（基线）",
         "图像 + terminal_data\n（端侧健康JSON，含当前页数据、异常标签）",
         "全部12维度",
         "评估模型在仅有端侧数据和图像时的综合能力，作为消融基线"),
        ("B（+description）",
         "A + description.json\n（图像场景文字描述，说明当前页是什么UI、展示什么数据）",
         "视觉识别准确率、数据引用质量、图像与上下文一致性",
         "验证图像场景描述是否能提升视觉理解和数据引用准确度"),
        ("C（+knowledge）",
         "A + knowledge.json\n（健康领域知识库：睡眠/心率/血氧各指标正常范围及临床意义）",
         "数值读取精度、数据引用质量、任务完成度",
         "验证健康知识注入是否能提升数值解读精度和回答完整度"),
        ("D（+both）",
         "A + description.json + knowledge.json",
         "数据引用质量、任务完成度",
         "验证两类辅助信息联合注入的协同效果"),
    ]
    for c, inp, dims, purpose in cond_data:
        row = cond_table.add_row().cells
        row[0].text = c
        row[1].text = inp
        row[2].text = dims
        row[3].text = purpose
    doc.add_paragraph("")

    # Image scene coverage
    doc.add_heading("测试图片场景覆盖", 1)
    p_img = doc.add_paragraph()
    p_img.add_run(
        "测试集共使用30张真实手机端健康App截图（PNG/JPG），覆盖7大健康场景类别，"
        "48个case由这30张图片与不同查询组合构成（部分图片出现在多个case中）。\n\n"
    )
    scene_table = doc.add_table(rows=1, cols=3)
    scene_table.style = "Table Grid"
    sh = scene_table.rows[0].cells
    sh[0].text = "场景类别"
    sh[1].text = "图片示例"
    sh[2].text = "主要评测内容"
    scene_data = [
        ("睡眠质量总览",
         "sleepScore_1/4.png, nightSleep_1.png, nightSleepGanttChart_1.png, naps_1.png",
         "睡眠得分识别、各阶段占比（深睡/浅睡/REM）、异常睡眠结构判断"),
        ("睡眠健康指标",
         "sleepApnea_1.png, sleepHeartrate_1.png, sleepSpo2_1.png",
         "呼吸暂停事件识别、睡眠心率/血氧趋势解读"),
        ("心率与HRV",
         "continueheartPage_3.png, silenceHeartPage_1.png, HRVPage_3.png",
         "静息心率异常识别、HRV趋势分析、连续心率图读数"),
        ("心律异常",
         "arrhythmiaStatisticsPage_1.png, todayArrhythmiaBarPage_1.png,\ntodayStatistics_1.png, sevenDayArrhythmiaBarPage_1.png/sevenDayStatistics_1.png",
         "房颤/室上速等心律失常事件统计、7日趋势、单日分布图读数"),
        ("血氧",
         "spo2MesultCompletePage_1.png, spo2TodayPage_1.png",
         "SpO2测量结果判读、日内血氧趋势、偏低阈值识别"),
        ("运动与活动",
         "exercisePage_1.png, recoverPage_1.png, todayHeadPage_2.png,\nactiveHeatPage_2.png, durationPage_2.png, weekDetailPage_2.jpg等",
         "运动记录解读、恢复评分、活动热力图、周运动时长统计"),
        ("健康总览仪表盘",
         "indexCard_1.png, scenarioSummary_1.png, domainCard_1.png",
         "综合健康评分卡、多维场景汇总、健康域评估卡片"),
    ]
    for cat, examples, content in scene_data:
        row = scene_table.add_row().cells
        row[0].text = cat
        row[1].text = examples
        row[2].text = content
    p_anomaly = doc.add_paragraph()
    p_anomaly.add_run("\n说明：").bold = True
    p_anomaly.add_run(
        '48个case中，36个case的terminal_data包含异常标签（"偏低"/"偏高"/"高于正常范围"等关键词），'
        "这36个case理论上应触发模型主动调用云侧健康数据API（get_health_data等工具）以获取更多参考数据。"
        "剩余12个case为正常指标，模型应直接基于端侧数据作答，不需要调用云侧工具。"
        '「工具调用分析」章节统计的"含异常标签的case数"即为该36个场景，用于评估模型的工具调用时机判断能力。'
    )
    doc.add_paragraph("")

    # Score summary table
    doc.add_heading("综合得分排名（条件A基线）", 1)
    rows = []
    for name in model_names:
        results_A = all_results[name].get("A", [])
        avgs = compute_dim_avgs(results_A, DIMS)
        overall = compute_overall(avgs)
        judged  = sum(1 for r in results_A if r.get("judgment"))
        rows.append((name, overall, judged, avgs))
    rows.sort(key=lambda x: -x[1])

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "排名 / 模型"
    hdr[1].text = "综合均分"
    hdr[2].text = "评分case数"
    hdr[3].text = "关键短板维度"
    for rank, (name, overall, judged, avgs) in enumerate(rows, 1):
        row = table.add_row().cells
        # Find weakest dim
        weakest = sorted([(d, v) for d, v in avgs.items() if v > 0], key=lambda x: x[1])
        weak_str = " | ".join(f"{d[:4]}:{v:.1f}" for d, v in weakest[:3])
        row[0].text = f"#{rank} {name}"
        row[1].text = f"{overall:.2f} / 10"
        row[2].text = str(judged)
        row[3].text = weak_str
    doc.add_paragraph("")

    # Case count notes
    p_note = doc.add_paragraph()
    p_note.add_run("评测说明：").bold = True
    p_note.add_run(
        "目标每模型每条件48 case。GLM-4.6V条件C因模型API不稳定有3次失败+2次token循环，"
        "实际有效case为43个；其余条件/模型均为47-48个（1 failed为正常概率性失败）。"
    )
    doc.add_paragraph("")

    # Charts
    doc.add_heading("评测图表", 1)
    for chart_file, caption in [
        ("bar_overall.png",    "图1: 5模型综合得分排名（条件A）"),
        ("heatmap_dims.png",   "图2: 12维度得分热图（条件A）"),
        ("radar_dims.png",     "图3: 多维能力雷达图（条件A）"),
        ("ablation_delta.png", "图4: 消融实验效果（description/knowledge贡献）"),
    ]:
        chart_path = charts_dir / chart_file
        if chart_path.exists():
            doc.add_paragraph(caption)
            doc.add_picture(str(chart_path), width=Inches(6))
            doc.add_paragraph("")

    # Per-model analysis
    doc.add_heading("逐模型分析", 1)
    for name, overall, judged, avgs in rows:
        doc.add_heading(f"{name}  (综合: {overall:.2f})", 2)
        # Dimension breakdown
        dim_rows = sorted(avgs.items(), key=lambda x: x[1])
        p = doc.add_paragraph()
        p.add_run("各维度得分：\n").bold = True
        for d, v in dim_rows:
            p.add_run(f"  {d}: {v:.2f}\n")

        # Ablation effect
        p2 = doc.add_paragraph()
        p2.add_run("消融实验增益：\n").bold = True
        for cond, label in [("B", "+description"), ("C", "+knowledge"), ("D", "+both")]:
            dims_cmp = ABLATION_DIMS.get(cond, [])
            results_cond = all_results[name].get(cond, [])
            avgs_cond = compute_dim_avgs(results_cond, dims_cmp)
            avgs_A_cmp = {d: avgs.get(d, 0) for d in dims_cmp}
            deltas = [(d, avgs_cond.get(d, 0) - avgs_A_cmp.get(d, 0))
                      for d in dims_cmp if avgs_A_cmp.get(d, 0) > 0 and avgs_cond.get(d, 0) > 0]
            if deltas:
                avg_delta = sum(v for _, v in deltas) / len(deltas)
                delta_str = " | ".join(f"{d[:4]}: {v:+.1f}" for d, v in deltas)
                p2.add_run(f"  [{cond}]{label}: avg delta={avg_delta:+.2f} ({delta_str})\n")

        # Notable cases
        results_A = all_results[name].get("A", [])
        tool_cases = [r for r in results_A if r.get("tool_calls") and r.get("judgment")]
        no_tool_anomaly = [r for r in results_A
                           if not r.get("tool_calls")
                           and any(kw in r.get("terminal_data", "")
                                   for kw in ["偏低", "偏高", "高于正常", "异常"])
                           and r.get("judgment")]
        if tool_cases:
            doc.add_paragraph(
                f"调用了工具的case: {len(tool_cases)} 条 | "
                f"有异常但未调工具: {len(no_tool_anomaly)} 条"
            )
        doc.add_paragraph("")

    # Tool call analysis
    doc.add_heading("工具调用分析", 1)
    p = doc.add_paragraph()
    p.add_run("各模型工具调用统计（条件A）：\n").bold = True
    for name in model_names:
        results_A = all_results[name].get("A", [])
        total = len([r for r in results_A if r.get("status") in ("done", "judged")])
        with_tools = len([r for r in results_A if r.get("tool_calls")])
        anomaly_cases = len([r for r in results_A
                             if any(kw in r.get("terminal_data", "")
                                    for kw in ["偏低", "偏高", "高于正常", "异常"])])
        all_tools = []
        for r in results_A:
            for tc in r.get("tool_calls", []):
                all_tools.append(tc.get("name", "?"))
        from collections import Counter
        top_tools = Counter(all_tools).most_common(5)
        p.add_run(f"\n  {name}: {with_tools}/{total} case有调用云工具"
                  f"（其中含异常标签、理应调用的case共{anomaly_cases}个）\n")
        if top_tools:
            p.add_run(f"    常用工具: {', '.join(f'{t}({c})' for t,c in top_tools)}\n")

    output.write_bytes(b"")  # ensure empty before save
    doc.save(str(output))
    print(f"[docx] {output}")


# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--run", help="指定 run 目录 (如 run_20260320_1000)")
    args = parser.parse_args()

    if args.run:
        results_dir = RESULTS_BASE / args.run
    else:
        runs = sorted(RESULTS_BASE.glob("run_*"), reverse=True)
        results_dir = runs[0] if runs else None

    if not results_dir or not results_dir.exists():
        print("[error] No results dir found.")
        return

    CHARTS_DIR.mkdir(parents=True, exist_ok=True)

    print(f"[report] Loading results from {results_dir}")
    all_results = load_results(results_dir)

    if not all_results:
        print("[error] No result files found (need *_results.json with judgment data)")
        return

    # Generate 4 charts
    chart_bar_overall(all_results, CHARTS_DIR / "bar_overall.png")
    chart_heatmap(    all_results, CHARTS_DIR / "heatmap_dims.png")
    chart_radar(      all_results, CHARTS_DIR / "radar_dims.png")
    chart_ablation(   all_results, CHARTS_DIR / "ablation_delta.png")

    # Generate DOCX
    generate_docx(all_results, CHARTS_DIR, EVAL_DIR / "summary_onapp.docx")

    print(f"\n[done] Charts → {CHARTS_DIR}")
    print(f"[done] Report → {EVAL_DIR / 'summary_onapp.docx'}")


if __name__ == "__main__":
    main()

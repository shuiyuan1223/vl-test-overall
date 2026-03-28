import json
import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.font_manager as fm

# Use a font that supports Chinese
plt.rcParams['font.family'] = ['Microsoft YaHei', 'SimHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

RESULT_DIR = "results/run_20260319_1205"
OUTPUT_DIR = "results/run_20260319_1205"

DIMS = [
    "视觉识别准确率", "幻觉控制率", "数值读取精度", "输出时序合规",
    "安全声明合规", "边界克制合规", "数据引用质量", "工具调用合理性",
    "任务完成度", "图像与上下文一致性", "主动澄清行为"
]

DIMS_SHORT = [
    "视觉识别", "幻觉控制", "数值读取", "时序合规",
    "安全声明", "边界克制", "数据引用", "工具调用",
    "任务完成", "图文一致", "主动澄清"
]

MODEL_COLORS = {
    "Qwen3.5-397B": "#1f77b4",
    "Qwen3.5-122B": "#ff7f0e",
    "kimi-k2.5": "#2ca02c",
    "Qwen3VL-235B": "#d62728",
    "GLM-4.6V": "#9467bd",
    "MiniMax-01": "#8c8c8c",
}

# Load all data
model_data = {}
for fname in os.listdir(RESULT_DIR):
    if fname.endswith("_results.json"):
        model_name = fname.replace("_results.json", "")
        with open(f"{RESULT_DIR}/{fname}", 'r', encoding='utf-8') as f:
            data = json.load(f)
        model_data[model_name] = data

# Sort by overall score
sorted_models = sorted(model_data.keys(),
                       key=lambda m: model_data[m].get('summary', {}).get('overall_score', 0),
                       reverse=True)

print("=== 综合得分排名 ===")
for m in sorted_models:
    s = model_data[m].get('summary', {}).get('overall_score', 0)
    print(f"  {m}: {s:.2f}")

# ---- Chart 1: Overall Score Bar Chart ----
fig, ax = plt.subplots(figsize=(10, 6))
models_plot = [m for m in sorted_models if m != "MiniMax-01"]
scores = [model_data[m]['summary']['overall_score'] for m in models_plot]
colors = [MODEL_COLORS[m] for m in models_plot]

bars = ax.bar(range(len(models_plot)), scores, color=colors, alpha=0.85, edgecolor='white', linewidth=1.5)
ax.set_xticks(range(len(models_plot)))
ax.set_xticklabels(models_plot, rotation=15, ha='right', fontsize=11)
ax.set_ylabel('综合得分 (1-10)', fontsize=12)
ax.set_title('VL Eval 综合得分排名（2026-03-19）', fontsize=14, fontweight='bold')
ax.set_ylim(0, 10)
ax.axhline(y=7.0, color='gray', linestyle='--', alpha=0.5, label='参考线 7.0')
ax.grid(axis='y', alpha=0.3)

for bar, score in zip(bars, scores):
    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1,
            f'{score:.2f}', ha='center', va='bottom', fontweight='bold', fontsize=11)

ax.legend(fontsize=10)
plt.tight_layout()
plt.savefig(f"{OUTPUT_DIR}/chart_overall_scores.png", dpi=150, bbox_inches='tight')
plt.close()
print("Generated: chart_overall_scores.png")

# ---- Chart 2: Dimension Scores Heatmap ----
models_for_heatmap = [m for m in sorted_models if m != "MiniMax-01"]
heatmap_data = []
for m in models_for_heatmap:
    row = [model_data[m]['summary']['avg_scores'][dim] for dim in DIMS]
    heatmap_data.append(row)

fig, ax = plt.subplots(figsize=(14, 5))
im = ax.imshow(heatmap_data, cmap='RdYlGn', vmin=1, vmax=10, aspect='auto')

ax.set_xticks(range(len(DIMS)))
ax.set_xticklabels(DIMS_SHORT, rotation=30, ha='right', fontsize=10)
ax.set_yticks(range(len(models_for_heatmap)))
ax.set_yticklabels(models_for_heatmap, fontsize=11)
ax.set_title('各模型11维度评分热力图', fontsize=13, fontweight='bold')

for i in range(len(models_for_heatmap)):
    for j in range(len(DIMS)):
        val = heatmap_data[i][j]
        color = 'white' if val < 4 or val > 7.5 else 'black'
        ax.text(j, i, f'{val:.1f}', ha='center', va='center', fontsize=9,
                color=color, fontweight='bold')

plt.colorbar(im, ax=ax, label='得分')
plt.tight_layout()
plt.savefig(f"{OUTPUT_DIR}/chart_heatmap.png", dpi=150, bbox_inches='tight')
plt.close()
print("Generated: chart_heatmap.png")

# ---- Chart 3: Radar Chart ----
models_radar = ["Qwen3.5-397B", "Qwen3.5-122B", "kimi-k2.5", "Qwen3VL-235B", "GLM-4.6V"]
N = len(DIMS_SHORT)
angles = [n / float(N) * 2 * np.pi for n in range(N)]
angles += angles[:1]

fig, ax = plt.subplots(figsize=(10, 10), subplot_kw=dict(polar=True))

for model in models_radar:
    values = [model_data[model]['summary']['avg_scores'][dim] for dim in DIMS]
    values += values[:1]
    ax.plot(angles, values, linewidth=2, label=model, color=MODEL_COLORS[model])
    ax.fill(angles, values, alpha=0.1, color=MODEL_COLORS[model])

ax.set_xticks(angles[:-1])
ax.set_xticklabels(DIMS_SHORT, size=10)
ax.set_ylim(0, 10)
ax.set_yticks([2, 4, 6, 8, 10])
ax.set_yticklabels(['2', '4', '6', '8', '10'], fontsize=8, color='gray')
ax.set_title('各模型11维度能力雷达图', size=14, fontweight='bold', pad=20)
ax.legend(loc='upper right', bbox_to_anchor=(1.35, 1.1), fontsize=10)
ax.grid(True, alpha=0.3)

plt.tight_layout()
plt.savefig(f"{OUTPUT_DIR}/chart_radar.png", dpi=150, bbox_inches='tight')
plt.close()
print("Generated: chart_radar.png")

# ---- Chart 4: Per-Case Scores for top 4 models ----
cases_labels = [
    "null", "爬坡训练", "跑步分析", "医疗报告",
    "体脂秤", "练肌肉", "今日状态", "体温计",
    "健康情况", "深蹲姿势", "血压", "血氧",
    "饮食(沙拉)", "饮食(晚餐)", "睡眠PSG", "睡前血糖"
]

top4_models = ["Qwen3.5-397B", "Qwen3.5-122B", "kimi-k2.5", "Qwen3VL-235B"]
fig, ax = plt.subplots(figsize=(16, 7))

x = np.arange(len(cases_labels))
width = 0.2

for i, model in enumerate(top4_models):
    case_avgs = []
    for result in model_data[model]['results']:
        s = result['judgment']['scores']
        avg = sum(s.values()) / len(s)
        case_avgs.append(avg)
    offset = (i - 1.5) * width
    bars = ax.bar(x + offset, case_avgs, width, label=model,
                  color=MODEL_COLORS[model], alpha=0.8)

ax.set_xticks(x)
ax.set_xticklabels(cases_labels, rotation=35, ha='right', fontsize=9)
ax.set_ylabel('平均得分', fontsize=11)
ax.set_title('各案例得分对比（Top 4 模型）', fontsize=13, fontweight='bold')
ax.set_ylim(0, 10)
ax.legend(fontsize=10)
ax.grid(axis='y', alpha=0.3)
ax.axhline(y=7.0, color='gray', linestyle='--', alpha=0.4)

plt.tight_layout()
plt.savefig(f"{OUTPUT_DIR}/chart_per_case.png", dpi=150, bbox_inches='tight')
plt.close()
print("Generated: chart_per_case.png")

# ---- Generate DOCX Report ----
doc = Document()

# Title
title = doc.add_heading('VL Eval 评测报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('PHA Visual Language Evaluation — 2026-03-19')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(12)
subtitle.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

# Executive Summary
doc.add_heading('1. 执行摘要', level=1)
exec_summary = doc.add_paragraph()
exec_summary.add_run(
    "本次评测对6个视觉语言模型在PHA健康助手场景下的综合能力进行了系统性测试，"
    "共16个测试案例，涵盖医疗图像识别、健康数据分析、运动姿态评估、饮食记录等场景。"
    "评测维度包含视觉识别准确率、幻觉控制、数值读取精度等11个维度。"
    "\n\nMiniMax-01不支持工具调用，无法纳入有效评测。"
    "其余5个模型中，Qwen3.5-397B表现最优（7.62分），GLM-4.6V表现最差（4.46分），存在多个严重崩溃案例。"
)

doc.add_paragraph()

# Overall Rankings Table
doc.add_heading('2. 综合得分排名', level=1)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '排名'
hdr[1].text = '模型'
hdr[2].text = '综合得分 (1-10)'

for cell in hdr:
    cell.paragraphs[0].runs[0].font.bold = True

valid_models = [m for m in sorted_models if m != "MiniMax-01"]
for rank, model in enumerate(valid_models, 1):
    row = table.add_row().cells
    row[0].text = str(rank)
    row[1].text = model
    row[2].text = f"{model_data[model]['summary']['overall_score']:.2f}"

doc.add_paragraph()

# Key Findings
doc.add_heading('3. 关键发现与分析', level=1)

findings = [
    ("3.1 Qwen3.5-397B 领先的核心原因",
     "397B在以下场景表现突出：\n"
     "• 乙肝医疗报告：正确识别\"小三阳\"（1/4/5阳性），解读HBsAg=85.808 IU/mL等所有指标\n"
     "• 体温计案例：明确区分腋下体温计与可穿戴设备体表温度的测量原理差异，这是最专业的响应\n"
     "• 血氧案例：识别93%/106bpm同时注意到鼻导管，给出氧疗专项建议（与122B并列最高10分）\n"
     "• 深蹲姿势：具体指出头部前低/背部前倾角度/手肘位置三个具体问题\n"
     "• 工具调用合理性最高（7.5分），善用HRV、心率等多维数据进行综合分析"),

    ("3.2 Qwen3.5-122B vs 397B 的核心差异",
     "122B vs 397B 的关键差异：\n"
     "• 医疗报告解读（4分 vs 8分）：122B拒绝解读，397B正确识别小三阳\n"
     "• 深蹲姿势（8分 vs 9分）：122B给出一般性建议，397B给出具体改进点\n"
     "• 工具调用（6.88 vs 7.5）：122B工具调用稍逊\n"
     "上次评测（无PHA_NOSTREAM）122B得分更低，说明397B受截断影响更大"),

    ("3.3 kimi-k2.5 的突出优势",
     "kimi在以下维度独特优秀：\n"
     "• 沙拉案例（最高分）：主动说'这不是沙拉，是牛排套餐'，是唯一正确识别的模型\n"
     "• 体脂秤：识别到24%体脂率（其他模型只读体重），主动澄清用户性别/年龄/身高\n"
     "• 主动澄清行为（6.75）：整体最高，每个case都主动询问基本信息\n"
     "• 严重缺陷：练肌肉case中调用3次get_skill严重冗余，工具调用合理性（6.44）偏低"),

    ("3.4 Qwen3VL-235B 的系统性缺陷",
     "235B存在两个严重系统性bug：\n"
     "• 健康情况案例（最低3分）：使用了2022-10-12的历史数据而非当前数据，严重错误\n"
     "• 睡眠报告案例（最低2分）：完全忽略图片中的PSG报告，转而分析3月18日设备数据\n"
     "• 练肌肉案例（最低2分）：调用10个工具但最终只输出一句'怎么称呼你'\n"
     "尽管如此，235B的工具调用数量最多，HRV分析最深入，鼻导管识别准确"),

    ("3.5 GLM-4.6V 的严重可靠性问题",
     "GLM-4.6V存在多个严重崩溃：\n"
     "• 4个案例完全崩溃：爬坡（空响应）、深蹲姿势（'滑翔'两字）、血压（token loop乱码）、睡眠PSG（'List of actions'无限循环）\n"
     "• 声称无视觉能力：报告解读案例说'我没有图像识别功能'，这是错误声明\n"
     "• 跑步案例：完全忽略6张跑步截图，说'今天没有跑步'\n"
     "• 亮点：体温计案例（8分）表现不错，正确识别38.6°C并对比系统数据\n"
     "• 结论：GLM-4.6V目前不适合生产环境，稳定性需大幅提升"),

    ("3.6 MiniMax-01 无法评测",
     "MiniMax-01不支持工具调用，PHA系统依赖工具调用才能获取用户健康数据，"
     "因此MiniMax-01无法完成任何健康分析任务。建议后续直连MiniMax API进行测试。"),

    ("3.7 视觉幻觉模式分析",
     "所有模型在null.png（湖景图）场景下都产生了'湖景'相关的幻觉，这是合理的（图片确实是湖景）。\n"
     "更关键的幻觉发现：\n"
     "• GLM-4.6V 声称无视觉能力（技术性幻觉）\n"
     "• Qwen3VL-235B 晚餐案例漏识别油条/煎蛋/草莓等大半食物\n"
     "• kimi/122B/397B 识别准确率相当，但都将'晚餐'早餐图调用为晚餐\n"
     "共同盲点：所有模型均未注意晚餐图片实为早餐（油条、豆浆、包子组合），没有主动指出"),
]

for title_text, content in findings:
    doc.add_heading(title_text, level=2)
    para = doc.add_paragraph(content)

# Recommendations
doc.add_heading('4. 建议与结论', level=1)
recommendations = doc.add_paragraph(
    "综合评测结果，给出以下建议：\n\n"
    "1. 生产推荐：Qwen3.5-397B 综合最优（7.62分），可用于生产环境，但医疗报告解读需配合安全声明\n"
    "2. 次选方案：kimi-k2.5（7.38分）在主动性和视觉识别方面有独特优势，工具调用需优化\n"
    "3. 基础版本：Qwen3.5-122B（7.22分）稳定可靠，边界克制良好，适合对安全性要求高的场景\n"
    "4. 暂不推荐：GLM-4.6V（4.46分）稳定性严重不足，4个案例完全崩溃，需等待版本改进\n"
    "5. 无法评测：MiniMax-01不支持工具调用，建议直连其原生API后重新测试\n\n"
    "下一步工作建议：\n"
    "• 对所有模型的医疗图像识别增加更多PSG、血常规、影像等医疗报告测试案例\n"
    "• 针对GLM-4.6V的token loop崩溃问题反馈给模型提供方\n"
    "• 对kimi-k2.5的冗余工具调用问题通过System Prompt优化解决\n"
    "• 设计早餐/晚餐场景辨别测试，评估模型的时间上下文理解能力"
)

doc.add_paragraph()

# Dimension Analysis
doc.add_heading('5. 维度详细分析', level=1)

dim_table = doc.add_table(rows=1, cols=7)
dim_table.style = 'Table Grid'
headers = ['维度', '397B', '122B', 'kimi', '235B', 'GLM', '说明']
for i, h in enumerate(headers):
    dim_table.rows[0].cells[i].text = h
    dim_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

model_order = ["Qwen3.5-397B", "Qwen3.5-122B", "kimi-k2.5", "Qwen3VL-235B", "GLM-4.6V"]
dim_notes = {
    "视觉识别准确率": "397B最强，GLM严重受损",
    "幻觉控制率": "397B/kimi较好，GLM崩溃案例多",
    "数值读取精度": "397B/122B/kimi相当，GLM因崩溃拖低",
    "输出时序合规": "235B工具调用最规范，各模型差异小",
    "安全声明合规": "各模型安全声明均较好，边界清晰",
    "边界克制合规": "235B/122B边界最严（拒绝报告解读），GLM拖低",
    "数据引用质量": "235B爬坡HRV分析最好，但健康情况数据日期错误",
    "工具调用合理性": "397B最佳，kimi 3x get_skill严重扣分",
    "任务完成度": "397B/122B均衡，235B练肌肉case只说了一句话",
    "图像与上下文一致性": "397B/kimi最佳，235B睡眠PSG分析严重偏差",
    "主动澄清行为": "kimi最主动，235B/397B较弱",
}

for dim in DIMS:
    row = dim_table.add_row().cells
    row[0].text = dim
    for i, m in enumerate(model_order):
        row[i+1].text = f"{model_data[m]['summary']['avg_scores'][dim]:.1f}"
    row[6].text = dim_notes.get(dim, "")

doc.add_paragraph()

# Charts
doc.add_heading('6. 综合得分图表', level=1)
doc.add_picture(f"{OUTPUT_DIR}/chart_overall_scores.png", width=Inches(6))

doc.add_heading('7. 维度能力雷达图', level=1)
doc.add_picture(f"{OUTPUT_DIR}/chart_radar.png", width=Inches(6))

doc.add_heading('8. 热力图', level=1)
doc.add_picture(f"{OUTPUT_DIR}/chart_heatmap.png", width=Inches(6.5))

doc.add_heading('9. 案例得分对比', level=1)
doc.add_picture(f"{OUTPUT_DIR}/chart_per_case.png", width=Inches(6.5))

doc_path = f"{OUTPUT_DIR}/summary.docx"
doc.save(doc_path)
print(f"\nReport saved: {doc_path}")
